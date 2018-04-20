from docx import Document
import re
import datetime


#Function for parsing date from string
#Функция для перевода даты в нужный формат
def parseDate(str):
    answer =""
    answer += str(d)
    return answer


def changeDateFormat(inputFile,outputFile):
    pDate = []
    document = Document('{0}.docx'.format(inputFile))
    for d in document.paragraphs:
        pDate.append(parseDate(d.text))
    doc = Document()
    for date in pDate:
        doc.add_paragraph(date)
    doc.save('{0}.docx'.format(outputFile))


changeDateFormat('some','test')